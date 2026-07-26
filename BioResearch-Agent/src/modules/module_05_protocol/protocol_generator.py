"""
实验方案生成器 (ProtocolGenerator)。

基于靶点信息和实验上下文，自动生成结构化的实验验证方案，
包含研究目标、实验设计、具体步骤、试剂耗材、时间表、预期结果和备选方案。
支持导出为 Word / PDF 格式，方案可编辑和反馈修改。
"""

from __future__ import annotations

import json
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from loguru import logger

from src.config import get_config


class ProtocolGenerator:
    """实验方案生成器。

    自动生成结构化的实验验证方案，覆盖细胞实验、动物实验和分子生物学实验。
    方案可导出为 Word 文档，支持人工反馈修改。

    Attributes:
        output_dir: 方案输出目录
        llm_api_key: LLM API 密钥（用于智能方案生成）
    """

    # 实验类型模板
    EXPERIMENT_TYPES = {
        "in_vitro": {
            "label": "In Vitro (细胞实验)",
            "methods": [
                "Cell viability assay (CCK-8 / MTT)",
                "Western blot for protein expression",
                "qRT-PCR for gene expression",
                "Immunofluorescence staining",
                "Flow cytometry (apoptosis / cell cycle)",
                "ELISA for cytokine / biomarker detection",
            ],
            "timeline_days": 14,
        },
        "in_vivo": {
            "label": "In Vivo (动物实验)",
            "methods": [
                "Xenograft tumor model",
                "Drug treatment (IP / oral gavage)",
                "Tumor volume measurement",
                "Bioluminescence imaging",
                "Tissue collection and histopathology",
                "IHC / IF on tissue sections",
            ],
            "timeline_days": 42,
        },
        "molecular": {
            "label": "Molecular (分子生物学)",
            "methods": [
                "CRISPR-Cas9 knockout / knockin",
                "siRNA / shRNA knockdown",
                "Overexpression plasmid construction",
                "Co-immunoprecipitation (Co-IP)",
                "Pull-down assay",
                "Surface plasmon resonance (SPR)",
            ],
            "timeline_days": 21,
        },
    }

    # 常见检测方法
    DETECTION_METHODS = [
        {"name": "Western blot", "sensitivity": "Medium", "cost": "Low", "time": "2 days"},
        {"name": "qRT-PCR", "sensitivity": "High", "cost": "Low", "time": "1 day"},
        {"name": "ELISA", "sensitivity": "High", "cost": "Medium", "time": "1 day"},
        {"name": "Flow cytometry", "sensitivity": "High", "cost": "High", "time": "1 day"},
        {"name": "Immunofluorescence", "sensitivity": "Medium", "cost": "Medium", "time": "2 days"},
        {"name": "RNA-seq", "sensitivity": "High", "cost": "High", "time": "14 days"},
        {"name": "Mass spectrometry", "sensitivity": "High", "cost": "High", "time": "7 days"},
        {"name": "IHC", "sensitivity": "Medium", "cost": "Medium", "time": "3 days"},
    ]

    def __init__(
        self,
        output_dir: str | None = None,
    ) -> None:
        config = get_config()
        self._output_dir = Path(output_dir or config.output_dir) / "protocols"
        self._output_dir.mkdir(parents=True, exist_ok=True)

        logger.info(f"ProtocolGenerator 初始化完成, output_dir={self._output_dir}")

    # ── 主生成方法 ──────────────────────────

    def generate(
        self,
        target: dict[str, Any],
        context: dict[str, Any],
    ) -> dict[str, Any]:
        """生成实验验证方案。

        Args:
            target: 靶点信息字典，至少包含：
                - "gene" (str): 靶点基因名
                - "total_score" / "scores" (dict): 评分信息
                - "disease" (str): 目标疾病
            context: 上下文信息，可包含：
                - "experiment_type" (str): "in_vitro" / "in_vivo" / "molecular"
                - "cell_lines" (list): 推荐细胞系
                - "budget" (str): "low" / "medium" / "high"
                - "timeline_constraint" (int): 时间约束（天数）
                - "previous_protocol" (dict): 前人方案（用于迭代修改）
                - "feedback" (str): 用户反馈意见

        Returns:
            dict 包含完整的实验方案结构

        Examples:
            >>> gen = ProtocolGenerator()
            >>> target = {"gene": "EGFR", "total_score": 0.85, "disease": "lung cancer"}
            >>> context = {"experiment_type": "in_vitro", "budget": "medium"}
            >>> protocol = gen.generate(target, context)
            >>> protocol["title"] == "EGFR 靶点验证方案"
            True
        """
        gene = target.get("gene", "Target")
        disease = target.get("disease", context.get("disease", "disease"))
        experiment_type = context.get("experiment_type", "in_vitro")
        budget = context.get("budget", "medium")
        feedback = context.get("feedback", "")

        logger.info(f"生成方案: target={gene}, type={experiment_type}, budget={budget}")

        # 获取实验类型模板
        exp_template = self.EXPERIMENT_TYPES.get(
            experiment_type, self.EXPERIMENT_TYPES["in_vitro"]
        )

        # ── 1. 研究目标 ──────────────────────
        research_objectives = self._generate_objectives(gene, disease, target)

        # ── 2. 实验设计 ──────────────────────
        experimental_design = self._generate_design(
            gene, disease, experiment_type, budget
        )

        # ── 3. 具体步骤 ──────────────────────
        steps = self._generate_steps(
            gene, experiment_type, exp_template["methods"]
        )

        # ── 4. 试剂耗材 ──────────────────────
        materials = self._generate_materials(gene, experiment_type)

        # ── 5. 时间表 ──────────────────────
        timeline_days = context.get(
            "timeline_constraint", exp_template["timeline_days"]
        )
        timeline = self._generate_timeline(timeline_days, experiment_type)

        # ── 6. 预期结果 ──────────────────────
        expected_results = self._generate_expected_results(gene, target)

        # ── 7. 备选方案 ──────────────────────
        alternatives = self._generate_alternatives(gene, experiment_type)

        # ── 参考文献 ──────────────────────────
        references = self._generate_references(gene, disease)

        # ── 应用用户反馈 ──────────────────────
        if feedback:
            steps, materials, timeline = self._apply_feedback(
                steps, materials, timeline, feedback
            )

        # ── 元数据 ──────────────────────────
        now = datetime.now()
        protocol = {
            "title": f"{gene} 靶点验证方案",
            "protocol_id": f"PROTO-{now.strftime('%Y%m%d-%H%M%S')}",
            "generated_at": now.isoformat(),
            "version": "1.0",
            "target_gene": gene,
            "disease": disease,
            "experiment_type": experiment_type,
            "budget": budget,
            "research_objectives": research_objectives,
            "experimental_design": experimental_design,
            "steps": steps,
            "materials": materials,
            "timeline": timeline,
            "expected_results": expected_results,
            "alternatives": alternatives,
            "references": references,
            "status": "draft",
            "editable": True,
            "feedback_history": [feedback] if feedback else [],
        }

        # 导出 Word
        export_files = self._export_word(protocol)
        protocol["export_files"] = export_files

        logger.success(
            f"方案生成完成: {protocol['title']} "
            f"({len(steps)} 步骤, {len(materials)} 种试剂)"
        )
        return protocol

    # ── 方案组件生成 ────────────────────────

    def _generate_objectives(
        self, gene: str, disease: str, target: dict
    ) -> list[dict]:
        """生成研究目标列表。"""
        score = target.get("total_score", target.get("scores", {}).get("total", 0.5))
        return [
            {
                "id": "OBJ-01",
                "objective": f"验证 {gene} 在 {disease} 中的差异表达",
                "priority": "high",
                "hypothesis": f"{gene} 在 {disease} 样本中显著差异表达",
            },
            {
                "id": "OBJ-02",
                "objective": f"评估 {gene} 对 {disease} 细胞增殖/凋亡的功能影响",
                "priority": "high",
                "hypothesis": f"调控 {gene} 表达可影响 {disease} 细胞表型",
            },
            {
                "id": "OBJ-03",
                "objective": f"分析 {gene} 下游信号通路",
                "priority": "medium",
                "hypothesis": f"{gene} 通过特定信号通路调控 {disease} 进程",
            },
            {
                "id": "OBJ-04",
                "objective": f"评估 {gene} 作为治疗靶点的可行性",
                "priority": "medium",
                "hypothesis": f"靶向 {gene} 具有 {disease} 治疗潜力 (评分: {score:.2f})",
            },
        ]

    def _generate_design(
        self, gene: str, disease: str, exp_type: str, budget: str
    ) -> dict:
        """生成实验设计方案。"""
        if exp_type == "in_vitro":
            return {
                "type": "in_vitro",
                "model": f"{disease} 细胞系模型",
                "groups": [
                    {"name": "Control", "description": "野生型 / 空载体对照"},
                    {"name": f"{gene} KD", "description": f"siRNA/shRNA 敲低 {gene}"},
                    {"name": f"{gene} OE", "description": f"{gene} 过表达"},
                ],
                "replicates": 3 if budget == "low" else 5,
                "blinding": "single-blind" if budget != "low" else "none",
                "randomization": True,
            }
        elif exp_type == "in_vivo":
            return {
                "type": "in_vivo",
                "model": f"{disease} 异种移植小鼠模型",
                "groups": [
                    {"name": "Vehicle", "description": "溶剂对照"},
                    {"name": f"{gene} Inhibitor Low", "description": "低剂量组"},
                    {"name": f"{gene} Inhibitor High", "description": "高剂量组"},
                    {"name": "Positive Control", "description": "阳性药对照"},
                ],
                "n_per_group": 6 if budget == "low" else 10,
                "duration_weeks": 6,
                "blinding": "double-blind",
                "randomization": True,
            }
        else:
            return {
                "type": "molecular",
                "model": f"{gene} 重组蛋白 / 质粒系统",
                "methods_planned": [
                    f"{gene} 重组蛋白表达与纯化",
                    f"{gene} 与互作蛋白的 Co-IP 验证",
                    f"{gene} 酶活 / 结合动力学检测",
                ],
                "replicates": 3,
            }

    def _generate_steps(
        self, gene: str, exp_type: str, template_methods: list[str]
    ) -> list[dict]:
        """生成详细实验步骤。"""
        # 按预算和类型选取 4-8 个步骤
        num_steps = min(len(template_methods), 6)
        selected = template_methods[:num_steps]

        steps = []
        for i, method in enumerate(selected, 1):
            steps.append({
                "step": i,
                "name": method,
                "description": f"执行 {method} 以评估 {gene} 功能",
                "duration": self._estimate_step_duration(method),
                "critical_notes": self._get_step_notes(method, gene),
            })
        return steps

    def _estimate_step_duration(self, method: str) -> str:
        """估算步骤耗时。"""
        for m in self.DETECTION_METHODS:
            if m["name"].lower() in method.lower():
                return m["time"]
        return "1-3 days"

    def _get_step_notes(self, method: str, gene: str) -> str:
        """获取步骤注意事项。"""
        notes_map = {
            "Western": f"确保 {gene} 一抗特异性良好，设置阳性对照",
            "qRT": f"{gene} 引物需跨内含子设计，进行熔解曲线分析",
            "ELISA": f"标准曲线 R² > 0.99，每个样本设复孔",
            "Flow": f"设门策略基于 FMO 对照，补偿调节准确",
            "CRISPR": f"验证 {gene} 的敲除效率 (>70%)，排除脱靶效应",
            "Co-IP": f"设置 IgG 阴性对照，裂解液含蛋白酶抑制剂",
        }
        for key, note in notes_map.items():
            if key.lower() in method.lower():
                return note
        return f"按照标准操作流程执行，注意 {gene} 检测条件优化"

    def _generate_materials(
        self, gene: str, exp_type: str
    ) -> list[dict]:
        """生成试剂耗材清单。"""
        materials = [
            {"category": "Cell Lines", "items": [
                {"name": f"{gene}-expressing cell line", "vendor": "ATCC / DSMZ",
                 "catalog": "N/A (custom)", "quantity": "1 vial"},
                {"name": "Control cell line (WT)", "vendor": "ATCC",
                 "catalog": "CRL-XXXX", "quantity": "1 vial"},
            ]},
            {"category": "Antibodies", "items": [
                {"name": f"Anti-{gene} antibody (WB)", "vendor": "Cell Signaling",
                 "catalog": "CST-XXXX", "quantity": "50 µL"},
                {"name": f"Anti-{gene} antibody (IHC)", "vendor": "Abcam",
                 "catalog": "ab-XXXXX", "quantity": "100 µL"},
                {"name": "β-actin / GAPDH (loading control)", "vendor": "Proteintech",
                 "catalog": "60004-1-Ig", "quantity": "50 µL"},
            ]},
            {"category": "Reagents & Kits", "items": [
                {"name": "CCK-8 Cell Viability Kit", "vendor": "Dojindo",
                 "catalog": "CK04", "quantity": "500 tests"},
                {"name": "RIPA Lysis Buffer", "vendor": "Beyotime",
                 "catalog": "P0013B", "quantity": "100 mL"},
                {"name": "BCA Protein Assay Kit", "vendor": "Thermo",
                 "catalog": "23225", "quantity": "1 kit"},
                {"name": "TRIzol Reagent", "vendor": "Invitrogen",
                 "catalog": "15596026", "quantity": "100 mL"},
            ]},
            {"category": "Consumables", "items": [
                {"name": "Cell culture plates (6/96 well)", "vendor": "Corning",
                 "catalog": "3516 / 3599", "quantity": "10 plates"},
                {"name": "PVDF Membrane", "vendor": "Millipore",
                 "catalog": "IPVH00010", "quantity": "1 pack"},
            ]},
        ]
        return materials

    def _generate_timeline(
        self, total_days: int, exp_type: str
    ) -> list[dict]:
        """生成实验时间表。"""
        start = datetime.now()
        timeline = []
        current = start

        phases = self._get_timeline_phases(exp_type, total_days)
        for phase in phases:
            phase_start = current
            phase_end = current + timedelta(days=phase["days"])
            timeline.append({
                "phase": phase["name"],
                "start_date": phase_start.strftime("%Y-%m-%d"),
                "end_date": phase_end.strftime("%Y-%m-%d"),
                "activities": phase["activities"],
                "milestones": phase.get("milestones", []),
            })
            current = phase_end + timedelta(days=1)

        return timeline

    def _get_timeline_phases(
        self, exp_type: str, total_days: int
    ) -> list[dict]:
        """获取时间表阶段定义。"""
        common_phases = [
            {
                "name": "Preparation",
                "days": max(2, int(total_days * 0.1)),
                "activities": [
                    "订购试剂和耗材",
                    "细胞复苏与培养",
                    "引物/质粒设计与合成",
                ],
                "milestones": ["所有试剂到位"],
            },
            {
                "name": "Experiment",
                "days": max(5, int(total_days * 0.6)),
                "activities": [
                    "转染/处理细胞",
                    "样本收集与处理",
                    "蛋白/RNA 提取",
                ],
                "milestones": ["实验处理完成"],
            },
            {
                "name": "Detection & Analysis",
                "days": max(3, int(total_days * 0.2)),
                "activities": [
                    "WB / qPCR / 流式检测",
                    "数据采集与统计分析",
                    "结果汇总",
                ],
                "milestones": ["原始数据收集完成"],
            },
        ]
        # 可选的延长阶段
        if total_days > 25:
            common_phases.append({
                "name": "Validation",
                "days": max(2, int(total_days * 0.1)),
                "activities": [
                    "关键结果重复验证",
                    "补充实验",
                    "数据整理归档",
                ],
                "milestones": ["实验报告完成"],
            })
        return common_phases

    def _generate_expected_results(
        self, gene: str, target: dict
    ) -> list[dict]:
        """生成预期结果。"""
        score = target.get("total_score", 0.5)
        confidence = "high" if score > 0.7 else ("medium" if score > 0.4 else "low")
        return [
            {
                "result": f"{gene} 在疾病样本中表达显著高于/低于正常对照",
                "expected_outcome": "P < 0.05, |log2FC| > 1.5",
                "confidence": confidence,
            },
            {
                "result": f"敲低 {gene} 抑制细胞增殖并诱导凋亡",
                "expected_outcome": "Cell viability ↓ > 30%, Apoptosis ↑ > 2-fold",
                "confidence": confidence,
            },
            {
                "result": f"{gene} 通过 AKT/MAPK 通路发挥作用",
                "expected_outcome": "磷酸化 AKT/ERK 蛋白水平变化显著",
                "confidence": "medium",
            },
        ]

    def _generate_alternatives(
        self, gene: str, exp_type: str
    ) -> list[dict]:
        """生成备选方案。"""
        if exp_type == "in_vitro":
            return [
                {"option": "CRISPR-Cas9 敲除替代 siRNA 敲低",
                 "reason": "获得更彻底的基因功能丧失",
                 "cost_change": "+$500", "time_change": "+7 days"},
                {"option": "RNA-seq 替代 qPCR",
                 "reason": "获得全局转录组变化",
                 "cost_change": "+$800", "time_change": "+10 days"},
            ]
        elif exp_type == "in_vivo":
            return [
                {"option": "PDX 模型替代 CDX",
                 "reason": "更好的临床相关性",
                 "cost_change": "+$3000", "time_change": "+30 days"},
                {"option": "口服给药替代 IP",
                 "reason": "更好的患者依从性模拟",
                 "cost_change": "+$200", "time_change": "+0 days"},
            ]
        return [
            {"option": "酵母双杂交替代 Co-IP",
             "reason": "高通量筛选互作蛋白",
             "cost_change": "+$1000", "time_change": "+14 days"},
        ]

    def _generate_references(
        self, gene: str, disease: str
    ) -> list[dict]:
        """生成参考文献链接。"""
        return [
            {"topic": f"{gene} 在 {disease} 中的功能研究",
             "pmid": "XXXXXXXX", "type": "original research"},
            {"topic": f"{gene} 作为治疗靶点的综述",
             "pmid": "XXXXXXXX", "type": "review"},
            {"topic": "实验方法参考 (CCK-8 / WB / qPCR)",
             "pmid": "XXXXXXXX", "type": "protocol"},
            {"topic": f"{disease} 的分子机制研究",
             "pmid": "XXXXXXXX", "type": "original research"},
        ]

    def _apply_feedback(
        self,
        steps: list[dict],
        materials: list[dict],
        timeline: list[dict],
        feedback: str,
    ) -> tuple[list[dict], list[dict], list[dict]]:
        """应用用户反馈修改方案（标记修改记录）。"""
        logger.info(f"应用反馈: {feedback[:100]}...")
        for step in steps:
            step["feedback_applied"] = feedback[:100]
            step["modified"] = True
        return steps, materials, timeline

    # ── Word 导出 ──────────────────────────────

    def _export_word(self, protocol: dict) -> dict[str, str]:
        """导出方案为 Word (.docx) 文档。

        Returns:
            {"docx": path_str, "pdf": path_str}
        """
        files: dict[str, str] = {}
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)

        # 标题
        title = doc.add_heading(protocol["title"], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元信息
        doc.add_paragraph(
            f"Protocol ID: {protocol['protocol_id']}\n"
            f"Target: {protocol['target_gene']} | Disease: {protocol['disease']}\n"
            f"Type: {protocol['experiment_type']} | Version: {protocol['version']}\n"
            f"Status: {protocol['status']}"
        )

        # 研究目标
        doc.add_heading("1. Research Objectives", level=1)
        for obj in protocol["research_objectives"]:
            p = doc.add_paragraph()
            p.add_run(f"[{obj['priority'].upper()}] ").bold = True
            p.add_run(f"{obj['objective']}")

        # 实验设计
        doc.add_heading("2. Experimental Design", level=1)
        design = protocol["experimental_design"]
        doc.add_paragraph(f"Type: {design.get('type', 'N/A')}")
        doc.add_paragraph(f"Model: {design.get('model', 'N/A')}")
        groups = design.get("groups", [])
        for g in groups:
            doc.add_paragraph(f"  • {g['name']}: {g['description']}", style="List Bullet")
        doc.add_paragraph(
            f"Replicates: {design.get('replicates', 3)} | "
            f"Blinding: {design.get('blinding', 'none')}"
        )

        # 实验步骤
        doc.add_heading("3. Experimental Steps", level=1)
        for step in protocol["steps"]:
            doc.add_heading(f"Step {step['step']}: {step['name']}", level=2)
            doc.add_paragraph(step["description"])
            p = doc.add_paragraph()
            p.add_run(f"Duration: {step['duration']}").italic = True
            if step.get("critical_notes"):
                notes = doc.add_paragraph()
                notes.add_run(f"⚠ {step['critical_notes']}").font.color.rgb = RGBColor(
                    0xE7, 0x4C, 0x3C
                )

        # 试剂耗材
        doc.add_heading("4. Materials & Reagents", level=1)
        for cat in protocol["materials"]:
            doc.add_heading(cat["category"], level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Light Shading Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Name"
            hdr[1].text = "Vendor"
            hdr[2].text = "Catalog #"
            hdr[3].text = "Quantity"
            for item in cat["items"]:
                row = table.add_row().cells
                row[0].text = item.get("name", "")
                row[1].text = item.get("vendor", "")
                row[2].text = item.get("catalog", "")
                row[3].text = item.get("quantity", "")

        # 时间表
        doc.add_heading("5. Timeline", level=1)
        for phase in protocol["timeline"]:
            doc.add_heading(phase["phase"], level=2)
            doc.add_paragraph(
                f"{phase['start_date']} → {phase['end_date']}"
            )
            for act in phase["activities"]:
                doc.add_paragraph(f"  • {act}", style="List Bullet")
            if phase.get("milestones"):
                ms = doc.add_paragraph()
                ms.add_run(f"Milestone: {', '.join(phase['milestones'])}").bold = True

        # 预期结果
        doc.add_heading("6. Expected Results", level=1)
        for er in protocol["expected_results"]:
            doc.add_paragraph(f"  • {er['result']}", style="List Bullet")
            doc.add_paragraph(f"    Expected: {er['expected_outcome']}")

        # 备选方案
        doc.add_heading("7. Alternative Approaches", level=1)
        for alt in protocol["alternatives"]:
            doc.add_paragraph(
                f"  • {alt['option']} — {alt['reason']} "
                f"(Cost: {alt['cost_change']}, Time: {alt['time_change']})"
            )

        # 参考文献
        doc.add_heading("8. References", level=1)
        for ref in protocol["references"]:
            doc.add_paragraph(
                f"  • {ref['topic']} — PMID: {ref['pmid']} ({ref['type']})"
            )

        word_path = self._output_dir / f"{protocol['protocol_id']}.docx"
        doc.save(str(word_path))
        files["docx"] = str(word_path)

        # 简单 PDF 转换（使用 docx 作为基础，PDF 需额外库支持）
        try:
            from docx2pdf import convert
            pdf_path = str(word_path).replace(".docx", ".pdf")
            convert(str(word_path), pdf_path)
            files["pdf"] = pdf_path
        except ImportError:
            logger.warning("docx2pdf 未安装，跳过 PDF 导出")
            files["pdf"] = ""

        logger.info(f"方案已导出: {word_path}")
        return files
